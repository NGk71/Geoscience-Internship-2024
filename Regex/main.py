import docx
import re
import spacy
import pandas as pd
from openpyxl import load_workbook

# Define keywords and their mapping
keyword_mapping = {
    "开始施工": ["开始施工"],
    "施工结束": ["施工结束"],
    "施工用时": ["施工用时"],
    "试压": ["试压"],
    "打备压": ["打备压"],
    "试挤用液": ["试挤用液"],
    "洗井用液": ["洗井用液"],
    "泵送桥塞用液": ["泵送桥塞用液"],
    "送球用液": ["送球用液"],
    "排空用液": ["排空用液"],
    "预处理酸": ["预处理酸"],
    "顶替液": ["顶替液"],
    "停泵反应": ["停泵反应"],
    "后打备压": ["后打备压"],
    "泵入前置液": ["泵入前置液"],
    "预搅拌浆体积": ["预搅拌浆体积"],
    "段塞加砂": ["段塞加砂"],
    "携砂液": ["携砂液"],
    "加砂": ["加砂"],
    "最高砂比": ["最高砂比"],
    "平均砂比": ["平均砂比"],
    "共加纤维": ["共加纤维"],
    "暂堵颗粒": ["暂堵颗粒"],
    "施工最高压力": ["施工最高压力"],
    "破裂压力": ["破裂压力"],
    "最高压力": ["最高压力"],
    "停泵油压": ["停泵油压"],
    "最大排量": ["最大排量"],
    "West": ["West"],
    "East": ["East"],
    "North": ["North"],
    "South": ["South"],
    "压裂方向": ["压裂方向"],
    "宽度": ["宽度"],
    "高度": ["高度"],
    "vol 1": ["vol 1"],
    "vol 2": ["vol 2"],
    "vol 3": ["vol 3"],
    "vol 4": ["vol 4"],
    "vol 5": ["vol 5"],
    "vol 6": ["vol 6"],
    "vol 7": ["vol 7"],
    "vol 8": ["vol 8"],
    "vol 9": ["vol 9"],
    "表面管处理压力": ["表面管处理压力"],
    "表面管压力": ["表面管压力"],
    "表面管初始激发压力": ["表面管初始激发压力"],
    "最大回流速率": ["最大回流速率"],
    "阶段顶部测深": ["阶段顶部测深"],
    "阶段底部测深": ["阶段底部测深"],
    "真垂直深度": ["真垂直深度"],
    "阶间分钟": ["阶间分钟"],
    "地层": ["地层"],
    "流体系统": ["流体系统"],
    "压裂系统": ["压裂系统"],
    "平均处理压力": ["平均处理压力"],
    "最大处理压力": ["最大处理压力"],
    "平均底压": ["平均底压"],
    "最大底压": ["最大底压"],
    "压裂梯度": ["压裂梯度"],
    "浆体积": ["浆体积"],
    "注入体积": ["注入体积"],
    "冲洗体积": ["冲洗体积"]
}
units = {
    "施工用时": "min",
    "试压": "MPa",
    "打备压": "MPa",
    "试挤用液": "m³",
    "洗井用液": "m³",
    "泵送桥塞用液": "m³",
    "送球用液": "m³",
    "排空用液": "m³",
    "预处理酸": "m³",
    "顶替液": "m³",
    "停泵反应": "min",
    "后打备压": "MPa",
    "泵入前置液": "m³",
    "段塞加砂": "m³",
    "携砂液": "m³",
    "加砂": "m³",
    "最高砂比": "%",
    "平均砂比": "%",
    "共加纤维": "Kg",
    "暂堵颗粒": "Kg",
    "施工最高压力": "MPa",
    "破裂压力": "MPa",
    "停泵油压": "MPa",
    "最大排量": "m³/min",
    "West": "m",
    "East": "m",
    "North": "m",
    "South": "m",
    "表面管处理压力": "MPa",
    "表面管压力": "MPa",
    "表面管初始激发压力": "MPa",
    "最大回流速率": "m³/min",
    "预搅拌浆体积": "m³",
    "宽度": "m",
    "高度": "m",
    "压裂方向": "",
    "vol 1": "m³",
    "vol 2": "m³",
    "vol 3": "m³",
    "vol 4": "m³",
    "vol 5": "m³",
    "vol 6": "m³",
    "vol 7": "m³",
    "vol 8": "m³",
    "vol 9": "m³",
    "阶段顶部测深": "ft",
    "阶段底部测深": "ft",
    "真垂直深度": "ft",
    "阶间分钟": "min",
    "地层": "",
    "流体系统": "",
    "压裂系统": "",
    "平均处理压力": "psi",
    "最大处理压力": "psi",
    "平均底压": "psi",
    "最大底压": "psi",
    "压裂梯度": "psi/ft",
    "浆体积": "bbl",
    "注入体积": "bbl",
    "冲洗体积": "bbl"
}
digit_to_chinese = {
    1: "一",
    2: "二",
    3: "三",
    4: "四",
    5: "五",
    6: "六",
    7: "七",
    8: "八",
    9: "九",
    10: "十",
    11: "十一",
    12: "十二",
    13: "十三",
    14: "十四",
    15: "十五",
    16: "十六",
    17: "十七",
    18: "十八",
    19: "十九",
    20: "二十",
    21: "二十一",
    22: "二十二",
    23: "二十三",
    24: "二十四",
    25: "二十五",
}
stage_mapping = {
    "stage1": "第一段",
    "stage2": "第二段",
    "stage3": "第三段",
    "stage4": "第四段",
    "stage5": "第五段",
    "stage6": "第六段",
    "stage7": "第七段",
    "stage8": "第八段",
    "stage9": "第九段",
    "stage10": "第十段",
    "stage11": "第十一段",
    "stage12": "第十二段",
    "stage13": "第十三段",
    "stage14": "第十四段",
    "stage15": "第十五段",
}
# Load the spaCy language model
nlp = spacy.load("en_core_web_sm")

# Preprocess text using spaCy
def preprocess_text(text):
    doc = nlp(text)
    processed_text = "\n".join([sent.text for sent in doc.sents])
    print(processed_text)
    return processed_text

# Extract well name from text
def extract_well_name(text):
    match = re.search(r'([\u4e00-\u9fa5\dA-Za-z-]+井)压裂施工总结', text)
    if match:
        print()
        return match.group(1)
    return None

# Extract text from DOCX file
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    start_scanning = True

    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if '施 工 单 位' in para.text:
            start_scanning = False

        if '五、施工工艺评价' in para.text:
            if i + 1 < len(doc.paragraphs) and '六、储层改造效果评价' not in doc.paragraphs[i + 1].text:
                start_scanning = True
                continue
        if '六、储层改造效果评价' in para.text:
            if i + 1 < len(doc.paragraphs) and '五、施工工艺评价' not in doc.paragraphs[i - 1].text:
                break
        if start_scanning:
            full_text.append(para.text)
            print("文本提取成功")
    return '\n'.join(full_text)

def search_keywords(text, keyword_mapping):
    results = []
    segments = re.split(r'(第[\d一二三四五六七八九十]+段)', text)
    segment_dict = {}

    for i in range(1, len(segments), 2):
        segment_name = segments[i].strip()
        piece = segments[i + 1]

        if segment_name not in segment_dict:
            segment_dict[segment_name] = {"段数": segment_name}

        for key, values in keyword_mapping.items():
            found_matches = set()  # Use a set to avoid duplicates
            for value in values:
                if key in ["开始施工", "施工结束"]:
                    pattern = rf"(\d{{1,2}}:\d{{2}})\D*{value}"
                else:
                    pattern = rf"{value}\s*[^\d]*\s*(\d+\.?\d*)"
                
                matches = re.findall(pattern, piece, re.DOTALL)
                found_matches.update(matches)

            found_matches = list(found_matches)  # Convert back to list to index
            if found_matches:
                for idx, match in enumerate(found_matches):
                    col_name = f"{key}{idx + 1}" if idx > 0 else key
                    segment_dict[segment_name][col_name] = match.strip()
            else:
                if key not in segment_dict[segment_name]:
                    segment_dict[segment_name][key] = None

    for segment_name in segment_dict:
        results.append(segment_dict[segment_name])
    print(results)
    return results

# Detect total segments in the tables
def detect_total_segments(tables):
    max_segment = 0
    pattern = re.compile(r'第(\d+|[一二三四五六七八九十十一十二十三十四十五十六十七十八十九二十]+)段')

    for table in tables:
        for row in table:
            joined_row = ''.join(row)
            match = pattern.search(joined_row)
            if match:
                segment_str = match.group(1)
                if segment_str.isdigit():
                    segment_num = int(segment_str)
                else:
                    segment_num = next(key for key, value in digit_to_chinese.items() if value == segment_str)
                max_segment = max(max_segment, segment_num)
    print("段数检测成功")
    return max_segment

# Extract tables from DOCX file
def extract_tables_from_docx(docx_path):
    doc = docx.Document(docx_path)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    print("表格提取成功")
    return tables

# Find tables for each segment


def find_segment_tables(tables):
    segment_tables = {}
    pattern = re.compile(r'第(\d+|[一二三四五六七八九十十一十二十三十四十五十六十七十八十九二十]+)段')

    current_segment = None
    for table in tables:
        for row in table:
            row_text = ''.join(row)
            match = pattern.search(row_text)
            if match:
                segment_str = match.group(1)
                if segment_str.isdigit():
                    segment_num = int(segment_str)
                else:
                    segment_num = next(key for key, value in digit_to_chinese.items() if value == segment_str)
                current_segment = segment_num
                if current_segment not in segment_tables:
                    segment_tables[current_segment] = []
            if current_segment:
                segment_tables[current_segment].append(row)
    print("段表格提取成功")
    return segment_tables

# Process the values to handle ranges by taking the average
def process_values(value):
    try:
        if '-' in str(value):
            start_val, end_val = map(float, value.split('-'))
            return (start_val + end_val) / 2
    except ValueError:
        pass
    return value

# Add new column based on conditions
# Add new column based on conditions
def only_table_needed(operations_df):
    if operations_df.empty:
        return operations_df
    
    # Identify the header row where "工序" appears in column A
    header_row_indices = operations_df[operations_df.iloc[:, 0].str.contains("工    序", na=False)].index
    
    if len(header_row_indices) == 0:
        return operations_df
    
    header_row_index = header_row_indices[0]
    
    # Use this row as the header row for the dataframe
    operations_df.columns = operations_df.iloc[header_row_index]
    operations_df = operations_df.drop(index=header_row_index).reset_index(drop=True)

    # Remove "/" in all cells and everything after it, and add one more column named "NO USE"
    operations_df = operations_df.apply(lambda x: x.str.split("/").str[0])
    operations_df["NO USE"] = ""

    if "备注" in operations_df.columns:
        operations_df.drop(columns=["备注"], inplace=True)

    return operations_df

# Save each segment's data to an Excel file
def save_segment_to_excel(segment_tables, segment_number, base_filename):
    if segment_number not in segment_tables:
        print(f"No operations data found for 第{digit_to_chinese[segment_number]}段. Skipping file creation.")
        return

    operations_data = segment_tables[segment_number]
    max_columns = max(len(row) for row in operations_data)
    formatted_data = [row + [""] * (max_columns - len(row)) for row in operations_data]

    # Apply processing to the data to handle ranges, except for "时间" column
    header = formatted_data[0]
    for i, row in enumerate(formatted_data[1:], start=1):
        for j, cell in enumerate(row):
            if header[j] != "时间":
                formatted_data[i][j] = process_values(cell)

    # Create a DataFrame to fill empty values
    operations_df = pd.DataFrame(formatted_data[1:], columns=formatted_data[0])

    # Add the new column based on specified conditions
    updated_df = only_table_needed(operations_df)

    # Fill empty values using forward fill method to connect lines
    updated_df.ffill(inplace=True)
    updated_df.bfill(inplace=True)

    filename = f"{base_filename}_process.xlsx"
    with pd.ExcelWriter(filename) as writer:
        updated_df.to_excel(writer, sheet_name=f'第{digit_to_chinese[segment_number]}段施工操作', index=False)
    print("段数据保存成功")
# Save results to an Excel file with customizable keyword names
def save_results_to_excel(results, well_name, base_filename, keyword_mapping):
    # Units associated with the keyword names


    filename = f"{base_filename}.xlsx"
    df = pd.DataFrame(results)
    
    # Generate the unit row dynamically based on the keyword_mapping names
    unit_row = {}
    for col in df.columns:
        base_name = col.split('1')[0]  # Extract base name before numeric suffix
        # Find the corresponding key in keyword_mapping
        unit = units.get(base_name, '')
        unit_row[col] = unit

    df = pd.concat([pd.DataFrame([unit_row]), df], ignore_index=True)
    df = df.dropna(subset=[col for col in df.columns if col != '段数'], how='all')

    df.insert(0, "井名", well_name)

    main_columns = ["井名", "段数"] + list(keyword_mapping.keys())

    organized_cols = []
    for col in main_columns:
        related_cols = [c for c in df.columns if c.startswith(col)]
        organized_cols.extend(related_cols)

    df = df[organized_cols]
    df.to_excel(filename, index=False)
    print("结果保存成功")
# Function to extract fracture data from the documents
def extract_fracture_data(docx_path):
    doc = docx.Document(docx_path)
    tables = extract_tables_from_docx(docx_path)
    data = []
    for table in tables:
        if table and "裂缝网络长(m)" in table[0]:
            for row in table[1:]:
                if len(row) < 7 or not row[1].isdigit():
                    continue
                stage_data = {
                    "Stage": row[0],
                    "West": row[1],
                    "East": row[2],
                    "North": row[1],
                    "South": row[2],
                    "压裂方向": row[6],
                    "宽度": row[4],
                    "高度": row[5]
                }
                data.append(stage_data)
    print("裂缝数据提取成功")
    return data
# Function to save the extracted data into the Excel file
def save_data_to_excel(data, excel_path):
    wb = load_workbook(excel_path)
    ws = wb.active

    # Define the column indexes based on headers
    headers = {cell.value: cell.column for cell in ws[1]}
    header_to_index = {
        "Stage": headers["段数"],
        "West": headers["West"],
        "East": headers["East"],
        "North": headers["North"],
        "South": headers["South"],
        "压裂方向": headers["压裂方向"],
        "宽度": headers["宽度"],
        "高度": headers["高度"]
    }
    # Insert data into the appropriate columns based on headers
    for stage_data in data:
        stage = stage_data.get("Stage")
        mapped_stage = stage_mapping.get(stage)
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            excel_stage = str(row[header_to_index["Stage"] - 1].value).strip()
            if excel_stage == mapped_stage:
                row[header_to_index["West"] - 1].value = stage_data.get("West")
                row[header_to_index["East"] - 1].value = stage_data.get("East")
                row[header_to_index["North"] - 1].value = stage_data.get("North")
                row[header_to_index["South"] - 1].value = stage_data.get("South")
                row[header_to_index["压裂方向"] - 1].value = stage_data.get("压裂方向")
                row[header_to_index["宽度"] - 1].value = stage_data.get("宽度")
                row[header_to_index["高度"] - 1].value = stage_data.get("高度")
    wb.save(excel_path)
    print("地震数据保存成功")
# Main script
def main():
    docx_path = 'testforall.docx'
    resultanalysis_path = 'resultanalysis.docx'
    excel_path = 'Well Operation Overview.xlsx'

    # Extract tables from the document
    tables = extract_tables_from_docx(docx_path)

    # Detect total number of segments
    total_segments = detect_total_segments(tables)

    # Find segment tables
    segment_tables = find_segment_tables(tables)

    # Save each segment's data to separate Excel files
    for i in range(1, total_segments + 1):
        base_filename = f'{i}'
        save_segment_to_excel(segment_tables, i, base_filename)

    # Extract and preprocess text
    text = extract_text_from_docx(docx_path)
    processed_text = preprocess_text(text)
    well_name = extract_well_name(processed_text)

    # Search keywords and save results to a combined Excel file
    results = search_keywords(processed_text, keyword_mapping)
    # Combine well overview data with the results
    combined_data = results
    save_results_to_excel(combined_data, well_name, 'Well Operation Overview', keyword_mapping)

    # Extract fracture data from both documents and save to Excel
    fracture_data_resultanalysis = extract_fracture_data(resultanalysis_path)
    fracture_data_testforall = extract_fracture_data(docx_path)
    combined_fracture_data = fracture_data_resultanalysis + fracture_data_testforall
    save_data_to_excel(combined_fracture_data, excel_path)

if __name__ == "__main__":
    main()



